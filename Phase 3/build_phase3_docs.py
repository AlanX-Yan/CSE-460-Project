from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
APP_DIR = ROOT / "SunDevilConnectApp"
PART1_DIR = ROOT / "Project Phase III - Part 1 " / "drafts"
PART2_DIR = ROOT / "Project Phase III - Part 2" / "drafts"
ASSET_DIR = ROOT / "Project Phase III - Part 1 " / "drafts" / "assets"


def font(size=22, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for item in candidates:
        try:
            return ImageFont.truetype(item, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_box(draw, xy, title, body, fill, outline="#344054"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 16), title, fill="#111827", font=font(22, True))
    y = y1 + 52
    for line in body:
        draw.text((x1 + 18, y), line, fill="#344054", font=font(17))
        y += 24


def arrow(draw, start, end, text=None):
    draw.line([start, end], fill="#475467", width=4)
    sx, sy = start
    ex, ey = end
    if ex >= sx:
        head = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    else:
        head = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    draw.polygon(head, fill="#475467")
    if text:
        mx, my = (sx + ex) / 2, (sy + ey) / 2
        draw.text((mx - 80, my - 26), text, fill="#475467", font=font(16))


def make_context_diagram(path):
    img = Image.new("RGB", (1400, 820), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((40, 32), "C4 Context Diagram - SunDevil Connect", fill="#111827", font=font(34, True))
    draw_box(draw, (500, 250, 900, 530), "SunDevil Connect", [
        "Web-based platform for discovering",
        "events, joining clubs, and managing",
        "club/admin workflows."
    ], "#fff7d6", "#8c1d40")
    draw_box(draw, (70, 130, 360, 310), "Student", [
        "Browses events",
        "Registers for events",
        "Requests club membership"
    ], "#e8f3ff")
    draw_box(draw, (70, 510, 360, 690), "Club Leader", [
        "Creates events",
        "Approves members",
        "Posts announcements"
    ], "#e2f0d9")
    draw_box(draw, (1040, 310, 1330, 500), "Administrator", [
        "Approves clubs",
        "Reviews flagged content",
        "Maintains platform quality"
    ], "#fce7f3")
    draw_box(draw, (1040, 590, 1330, 745), "Browser Storage", [
        "Prototype seed data",
        "In-memory application state"
    ], "#edf2f7")
    arrow(draw, (360, 220), (500, 330), "uses")
    arrow(draw, (360, 600), (500, 450), "manages")
    arrow(draw, (1040, 405), (900, 405), "oversees")
    arrow(draw, (900, 530), (1040, 650), "reads/writes")
    img.save(path)


def make_component_diagram(path):
    img = Image.new("RGB", (1500, 920), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "C4 Component Diagram - Browser Application", fill="#111827", font=font(34, True))
    draw_box(draw, (40, 150, 330, 340), "Boundary/UI", [
        "Student view",
        "Leader dashboard",
        "Admin panel"
    ], "#e8f3ff")
    draw_box(draw, (440, 150, 760, 340), "SunDevilConnectFacade", [
        "Simplified entry point",
        "Coordinates user actions",
        "Hides controller details"
    ], "#fff7d6", "#8c1d40")
    draw_box(draw, (900, 90, 1210, 230), "AuthController", ["Role selection", "Current user lookup"], "#ffffff")
    draw_box(draw, (900, 270, 1210, 430), "EventController", ["Search/filter events", "Register/cancel", "Create events"], "#ffffff")
    draw_box(draw, (900, 470, 1210, 630), "ClubController", ["Membership requests", "Approvals", "Announcements"], "#ffffff")
    draw_box(draw, (900, 670, 1210, 830), "AdminController", ["Club approvals", "Flagged content review"], "#ffffff")
    draw_box(draw, (1260, 190, 1460, 360), "State Pattern", ["EventState", "Draft/Published", "Full/Cancelled"], "#eadcf8")
    draw_box(draw, (1260, 510, 1460, 710), "Command Pattern", ["DashboardCommand", "Create event", "Approve member", "Review content"], "#e2f0d9")
    draw_box(draw, (440, 540, 760, 760), "Entities + Store", [
        "User, Student, Club",
        "Event, Membership",
        "Announcement, FlaggedContent"
    ], "#edf2f7")
    arrow(draw, (330, 245), (440, 245), "calls")
    arrow(draw, (760, 245), (900, 170), "")
    arrow(draw, (760, 255), (900, 350), "")
    arrow(draw, (760, 265), (900, 550), "")
    arrow(draw, (760, 275), (900, 750), "")
    arrow(draw, (1060, 430), (1260, 275), "uses")
    arrow(draw, (1060, 630), (1260, 605), "executes")
    arrow(draw, (1060, 830), (1260, 625), "executes")
    arrow(draw, (900, 350), (760, 610), "updates")
    arrow(draw, (900, 550), (760, 650), "updates")
    arrow(draw, (900, 750), (760, 700), "updates")
    img.save(path)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(140, 29, 64)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(102, 112, 133)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_part1():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    context_path = ASSET_DIR / "c4-context.png"
    component_path = ASSET_DIR / "c4-component.png"
    make_context_diagram(context_path)
    make_component_diagram(component_path)

    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "Class Project Phase III - Part 1 Submission", "SunDevil Connect | Group 6 - Individual Contribution | Xu Yan")

    doc.add_heading("Architecture Explanation and Architecture Diagram [15 Points]", level=1)
    doc.add_heading("Selected Architecture: Layered Client-Side Architecture", level=2)
    doc.add_paragraph(
        "For the SunDevil Connect implementation, I selected a layered client-side architecture. "
        "The application is organized into a presentation layer, a coordination layer, a business logic layer, and a domain data layer. "
        "The presentation layer contains the Student view, Club Leader dashboard, and Admin panel. These screens collect input and display event, club, membership, and moderation information. "
        "The coordination layer is represented by the SunDevilConnectFacade, which gives the UI a simple set of methods for browsing events, registering for events, joining clubs, and executing dashboard commands. "
        "The business logic layer contains AuthController, EventController, ClubController, and AdminController. These controllers enforce the main rules of the system. "
        "The domain layer contains the main entities, including User, Student, Club, Event, ClubMembership, EventRegistration, Announcement, and FlaggedContent."
    )
    doc.add_paragraph(
        "This architecture is a good fit for the current implementation because the project is a prototype that must demonstrate the complete workflow without requiring a deployed backend. "
        "It still preserves separation of concerns, so the system does not become one large script. The facade keeps the UI independent from controller details. "
        "The controllers keep business rules away from the HTML. The entity objects and data store make the application state easy to understand and modify. "
        "The implementation also applies two Phase II patterns directly: the State Pattern controls event lifecycle behavior, and the Command Pattern handles leader/admin actions such as creating events, approving memberships, posting announcements, and reviewing flagged content."
    )

    doc.add_heading("Comparison With Alternative Architecture", level=2)
    doc.add_paragraph(
        "One alternative would be a simple monolithic front-end script where every button directly changes global arrays and page elements. "
        "That approach would be faster for a very small demo, but it would become difficult to maintain as soon as more roles and workflows are added. "
        "For example, student registration rules, club approval rules, event capacity rules, and flagged content moderation would all become mixed with DOM manipulation. "
        "Another alternative would be a microservices architecture with separate services for users, events, clubs, and moderation. "
        "Microservices could support a large production system, but they would add unnecessary deployment and integration complexity for this course increment."
    )
    doc.add_paragraph(
        "The layered client-side architecture is better for this phase because it is realistic enough to show solid design while still being feasible for an individual implementation. "
        "It supports incremental growth: the current in-memory data store could later be replaced with a database-backed API while keeping the same UI and facade methods. "
        "This makes the architecture maintainable, testable, and appropriate for the Phase III implementation schedule."
    )

    doc.add_heading("C4 Context Diagram", level=2)
    doc.add_picture(str(context_path), width=Inches(6.3))
    doc.add_paragraph(
        "The context diagram shows the three main external actors: Student, Club Leader, and Administrator. "
        "Each actor interacts with SunDevil Connect through the browser-based application. The prototype stores seeded data and user actions in browser application state."
    )

    doc.add_heading("C4 Component Diagram", level=2)
    doc.add_picture(str(component_path), width=Inches(6.3))
    doc.add_paragraph(
        "The component diagram shows how UI pages communicate with the SunDevilConnectFacade. "
        "The facade coordinates the controllers, and the controllers update domain entities and application state. "
        "EventState classes handle event status behavior, while DashboardCommand classes encapsulate dashboard operations."
    )

    doc.add_heading("Phase III Part 1 Code [10 Points]", level=1)
    doc.add_paragraph("Code folder prepared for upload:")
    add_bullets(doc, [str(APP_DIR)])
    doc.add_paragraph("Google Drive link to paste after upload: [PASTE PHASE III PART 1 CODE LINK HERE]")

    doc.add_heading("Video Link [5 Points]", level=1)
    doc.add_paragraph(
        "Prepare a 3-5 minute video explaining the architecture and demonstrating the student-side implementation: event browsing, filtering, event registration, and club membership request."
    )
    doc.add_paragraph("Video link to paste after upload: [PASTE PHASE III PART 1 VIDEO LINK HERE]")

    path = PART1_DIR / "PhaseIII_Part1_Submission_Draft.docx"
    doc.save(path)
    return path


def build_part2():
    doc = Document()
    set_doc_styles(doc)
    add_title(doc, "Class Project Phase III - Part 2 Submission", "SunDevil Connect | Group 6 - Individual Contribution | Xu Yan")

    doc.add_heading("Phase III Part 2 Code [5 Points]", level=1)
    doc.add_paragraph(
        "Part 2 extends the implementation with Club Leader and Administrator workflows. "
        "The Club Leader dashboard supports creating events, approving pending club membership requests, and posting announcements. "
        "The Admin panel supports approving pending clubs and reviewing flagged content. These actions are implemented using command objects, including CreateEventCommand, ApproveMemberCommand, PostAnnouncementCommand, and ReviewFlaggedContentCommand."
    )
    doc.add_paragraph("Code folder prepared for upload:")
    add_bullets(doc, [str(APP_DIR)])
    doc.add_paragraph("Google Drive link to paste after upload: [PASTE PHASE III PART 2 CODE LINK HERE]")

    doc.add_heading("Video Link [5 Points]", level=1)
    doc.add_paragraph(
        "Prepare a 3-5 minute video demonstrating the Part 2 features: create a leader event, approve a membership request, post an announcement, approve a club, and mark flagged content as removed or dismissed."
    )
    doc.add_paragraph("Video link to paste after upload: [PASTE PHASE III PART 2 VIDEO LINK HERE]")

    path = PART2_DIR / "PhaseIII_Part2_Submission_Draft.docx"
    doc.save(path)
    return path


def write_demo_script():
    script = PART2_DIR / "PhaseIII_Demo_Script.md"
    script.write_text(
        """# Phase III Demo Script

## Part I Demo

1. Open `SunDevilConnectApp/index.html`.
2. Show the Student tab.
3. Search for an event and filter by category or cost.
4. Register for `AI Project Night`.
5. Show that `My Activity` updates with the registration.
6. Join a club and show that the membership request appears as pending.
7. Briefly explain that UI actions go through `SunDevilConnectFacade`, then controllers update entities.

## Part II Demo

1. Switch to the Club Leader tab.
2. Create a new event using the form.
3. Approve the pending membership request.
4. Post a club announcement.
5. Switch to the Admin tab.
6. Approve the pending club.
7. Mark one flagged content item as removed or dismissed.
8. Explain that these dashboard actions use Command Pattern classes.
""",
        encoding="utf-8",
    )
    return script


if __name__ == "__main__":
    print(build_part1())
    print(build_part2())
    print(write_demo_script())
